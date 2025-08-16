from openpyxl import load_workbook
from docx import Document
import os
from docx.shared import Pt
import dashscope
from dashscope import Generation
import streamlit as st
from pathlib import Path
import pandas as pd
from io import BytesIO
import tempfile

# 设置页面布局
st.set_page_config(page_title="智能尽职调查报告生成系统", layout="wide")
st.title("📑 智能尽职调查报告生成系统")


def multi_cell_to_word_template(excel_path, word_template_path, output_path, cell_mapping):
    """
    从Excel提取多个单元格内容并插入到Word模板的多个位置
    """
    try:
        # 1. 读取Excel数据
        wb = load_workbook(excel_path)
        ws = wb.active

        # 2. 准备替换数据字典
        replace_data = {}
        for placeholder, cell_ref in cell_mapping.items():
            col_letter = ''.join([c for c in cell_ref if c.isalpha()])
            row_num = int(''.join([c for c in cell_ref if c.isdigit()]))

            # Excel列字母转数字
            col_num = sum((ord(c.upper()) - ord('A') + 1) * (26 ** i)
                          for i, c in enumerate(reversed(col_letter)))

            cell_value = ws.cell(row=row_num, column=col_num).value
            replace_data[placeholder] = str(cell_value) if cell_value is not None else ""

        # 3. 处理Word模板
        doc = Document(word_template_path)

        # 替换内容
        for paragraph in doc.paragraphs:
            for placeholder, value in replace_data.items():
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, value)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for placeholder, value in replace_data.items():
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, value)

        # 设置文档样式
        style = doc.styles['Normal']
        style.font.name = '仿宋_GB2312'
        style.font.size = Pt(16)

        # 4. 保存结果
        doc.save(output_path)
        return True

    except Exception as e:
        st.error(f"生成Word文档时出错: {str(e)}")
        return False


def extract_word_content(file_path):
    """安全读取Word内容"""
    try:
        doc = Document(file_path)
        return "\n".join(para.text for para in doc.paragraphs)
    except Exception as e:
        st.error(f"读取Word内容失败: {str(e)}")
        return ""


def call_qwen_api(prompt):
    """调用通义千问API"""
    try:
        dashscope.api_key = os.getenv('DASHSCOPE_KEY', 'sk-ff22dbef6b2a4596b52b334df5b77ec3')

        response = Generation.call(
            model="qwen-turbo",
            prompt=prompt
        )

        if response.status_code == 200:
            return response.output['text']
        else:
            raise Exception(response.message)
    except Exception as e:
        st.error(f"API调用失败: {str(e)}")
        return ""


def download_word_doc(file_path):
    """下载Word文档"""
    try:
        with open(file_path, "rb") as f:
            st.download_button(
                label="📥 下载尽职调查报告",
                data=f,
                file_name=Path(file_path).name,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    except Exception as e:
        st.error(f"下载文件失败: {str(e)}")


def main():
    # 初始化目录
    output_dir = Path("outputs")
    output_dir.mkdir(exist_ok=True)

    # 文件上传
    st.sidebar.header("1. 上传文件")
    excel_file = st.sidebar.file_uploader(
        "选择Excel文件",
        type=["xlsx", "xls"],
        help="请上传包含尽职调查数据的Excel文件"
    )

    # 模板路径设置
    template_path = Path(r"C:\Users\Administrator\Desktop\新建文件\附件1.尽职调查报告模板.docx")
    if not template_path.exists():
        st.error(f"模板文件不存在: {template_path}")
        return

    # 单元格映射关系
    cell_mapping = {
        '{{A}}': 'D2', '{{B}}': 'D3', '{{C}}': 'D7', '{{D}}': 'C7',
        '{{E}}': 'D8', '{{F}}': 'C8', '{{G}}': 'D9', '{{H}}': 'C9',
        '{{I}}': 'D10', '{{J}}': 'C10', '{{K}}': 'D11', '{{L}}': 'C11',
        '{{M}}': 'D12', '{{N}}': 'C12', '{{O}}': 'D13', '{{P}}': 'C13',
        '{{Q}}': 'D14', '{{R}}': 'C14'
    }

    if excel_file:
        # 显示上传信息
        st.success(f"已上传文件: {excel_file.name}")

        # 使用临时文件处理
        with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as tmp:
            tmp.write(excel_file.getbuffer())
            excel_path = tmp.name

        # 生成报告
        output_file = output_dir / "surveyreport.docx"

        with st.spinner("正在生成尽职调查报告..."):
            if multi_cell_to_word_template(excel_path, str(template_path), str(output_file), cell_mapping):
                st.success("报告生成成功！")

                # 提取关键信息
                wb = load_workbook(excel_path)
                sheet = wb.active
                time1 = sheet['D5'].value
                time2 = sheet['D6'].value
                bank = sheet['D2'].value
                name = sheet['D3'].value

                # 生成AI结论
                word_content = extract_word_content(str(output_file))
                prompt = f"""根据以下内容生成风险评估结论：
                支行名称: {bank}
                负责人: {name}
                任期: {time1} 至 {time2}
                ------
                {word_content}
                """

                with st.spinner("AI正在生成风险评估结论..."):
                    ai_result = call_qwen_api(prompt)
                    if ai_result:
                        # 将AI结果添加到报告
                        doc = Document(str(output_file))
                        doc.add_heading("AI风险评估结论", level=1)
                        doc.add_paragraph(ai_result)
                        final_output = output_dir / "final_report.docx"
                        doc.save(str(final_output))

                        # 显示结果
                        st.subheader("AI生成结论")
                        st.text_area("结论内容", ai_result, height=300)

                        # 提供下载
                        download_word_doc(str(final_output))
            else:
                st.error("报告生成失败，请检查模板和输入数据")

        # 清理临时文件
        try:
            os.unlink(excel_path)
        except:
            pass


if __name__ == "__main__":
    main()